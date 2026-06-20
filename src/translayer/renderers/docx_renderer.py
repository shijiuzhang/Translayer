"""DOCX renderer — lossless write-back of translations into the original file.

Reopens the *original* docx to preserve all formatting, locates each
paragraph/cell via the block's SourceRef, and replaces the text while
keeping the first run's formatting.
"""

from __future__ import annotations

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from translayer.ir.models import Block, DocumentIR
from translayer.plugins import registry


class DocxRenderer:
    name = "docx"

    def supported_formats(self) -> list[str]:
        return ["docx"]

    def render(self, ir: DocumentIR, input_path: str, output_path: str) -> None:
        doc = Document(input_path)

        paragraphs, table_cells = self._build_indices(doc)

        for block in ir.blocks:
            if block.type == "image":
                self._render_image(block, ir, doc)
                continue
            if block.target_text is None:
                continue
            ref = block.source_ref
            if ref.kind == "table_cell" and ref.row is not None and ref.col is not None:
                para = table_cells.get((ref.paragraph_index, ref.row, ref.col))
                if para:
                    self._set_paragraph_text(para, block.target_text)
            else:
                para = paragraphs.get(ref.paragraph_index)
                if para:
                    self._set_paragraph_text(para, block.target_text)

        doc.save(output_path)

    def _build_indices(self, doc: Document) -> tuple[dict[int, Paragraph], dict[tuple[int, int, int], Paragraph]]:
        """Build unified paragraph and table-cell indices matching the parser's global counter."""
        paragraphs: dict[int, Paragraph] = {}
        table_cells: dict[tuple[int, int, int], Paragraph] = {}
        counter = 0
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if tag == "p":
                paragraphs[counter] = Paragraph(element, doc)
                counter += 1
            elif tag == "tbl":
                table = Table(element, doc)
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        for para in cell.paragraphs:
                            table_cells[(counter, r_idx, c_idx)] = para
                            counter += 1
        return paragraphs, table_cells

    @staticmethod
    def _set_paragraph_text(para: Paragraph, text: str) -> None:
        runs = para.runs
        if not runs:
            run = para.add_run()
            run.text = text
            return
        runs[0].text = text
        for extra in runs[1:]:
            extra.text = ""

    def _render_image(self, block: Block, ir: DocumentIR, doc: Document) -> None:
        ref = block.source_ref
        if ref.image_id is None:
            return
        image = ir.image_by_id(ref.image_id)
        if image is None or not image.localized_data_ref:
            return
        try:
            with open(image.localized_data_ref, "rb") as fh:
                new_blob = fh.read()
        except OSError:
            return

        for rel in doc.part.rels.values():
            if "image" in rel.reltype and ref.image_id in (rel.rId or ""):
                try:
                    rel.target_part._blob = new_blob
                except AttributeError:
                    pass
                break


registry.register("renderer", "docx")(DocxRenderer)
