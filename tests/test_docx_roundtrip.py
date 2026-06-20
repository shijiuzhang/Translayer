"""DOCX parser + renderer round-trip tests."""

from __future__ import annotations

from docx import Document

from translayer.parsers.docx_parser import DocxParser
from translayer.renderers.docx_renderer import DocxRenderer


def _make_sample_docx(tmp_path) -> str:
    doc = Document()
    doc.add_heading("Quarterly Report", level=1)
    doc.add_paragraph("Revenue grew strongly this quarter.")
    doc.add_paragraph("Costs were well managed across all departments.")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Revenue"
    table.cell(1, 1).text = "5000"

    doc.add_paragraph("Next quarter looks promising.", style="List Bullet")

    out = str(tmp_path / "sample.docx")
    doc.save(out)
    return out


def test_parse_extracts_paragraphs_and_headings(tmp_path):
    sample = _make_sample_docx(tmp_path)
    ir = DocxParser().parse(sample, {"source_lang": "en", "target_lang": "zh"})

    texts = [b.source_text for b in ir.blocks if b.type != "image"]
    assert "Quarterly Report" in texts
    assert "Revenue grew strongly this quarter." in texts
    assert "Costs were well managed across all departments." in texts

    headings = [b for b in ir.blocks if b.type == "heading"]
    assert len(headings) >= 1
    assert headings[0].semantic_role == "h1"


def test_parse_extracts_table_cells(tmp_path):
    sample = _make_sample_docx(tmp_path)
    ir = DocxParser().parse(sample, {"source_lang": "en", "target_lang": "zh"})

    table_blocks = [b for b in ir.blocks if b.type == "table_cell"]
    assert len(table_blocks) == 4
    texts = {b.source_text for b in table_blocks}
    assert "Metric" in texts
    assert "Revenue" in texts
    assert "5000" in texts


def test_parse_extracts_list_items(tmp_path):
    sample = _make_sample_docx(tmp_path)
    ir = DocxParser().parse(sample, {"source_lang": "en", "target_lang": "zh"})

    list_items = [b for b in ir.blocks if b.type == "list_item"]
    assert len(list_items) >= 1
    assert "Next quarter looks promising." in [b.source_text for b in list_items]


def test_render_writes_back_translations(tmp_path):
    sample = _make_sample_docx(tmp_path)
    parser, renderer = DocxParser(), DocxRenderer()
    ir = parser.parse(sample, {})

    mapping = {
        "Quarterly Report": "季度报告",
        "Revenue grew strongly this quarter.": "本季度收入大幅增长。",
        "Costs were well managed across all departments.": "各部门成本控制良好。",
        "Metric": "指标",
        "Value": "数值",
        "Revenue": "收入",
        "5000": "5000",
    }
    for b in ir.blocks:
        if b.source_text in mapping:
            b.target_text = mapping[b.source_text]

    out = str(tmp_path / "out.docx")
    renderer.render(ir, sample, out)

    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for cell_text in ["指标", "数值", "收入"]:
        found = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell_text in cell.text:
                        found = True
        assert found, f"{cell_text} not found in tables"

    assert "季度报告" in full_text
    assert "本季度收入大幅增长。" in full_text
    assert "Quarterly Report" not in full_text


def test_roundtrip_preserves_structure(tmp_path):
    sample = _make_sample_docx(tmp_path)
    parser, renderer = DocxParser(), DocxRenderer()
    ir = parser.parse(sample, {})

    out = str(tmp_path / "identity.docx")
    renderer.render(ir, sample, out)

    src, dst = Document(sample), Document(out)
    assert len(src.paragraphs) == len(dst.paragraphs)
    assert len(src.tables) == len(dst.tables)


def test_source_refs_have_paragraph_index(tmp_path):
    sample = _make_sample_docx(tmp_path)
    ir = DocxParser().parse(sample, {})

    for b in ir.blocks:
        if b.type != "image":
            assert b.source_ref.paragraph_index is not None
